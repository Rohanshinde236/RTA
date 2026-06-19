"""
Generates documents/RTA_Manual_vs_Automation.docx using python-docx
(natively Word-compatible — avoids docx-js field/schema issues).

Usage:  python make_process_doc.py [output_path]
"""
import sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)
DARK = RGBColor(0x22, 0x22, 0x22)
LIGHT = 'EAF1F8'
HEADFILL = '1F3A5F'

doc = Document()

# ── base styling ──────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK

for lvl, sz, col in [(1, 16, NAVY), (2, 13, BLUE), (3, 12, DARK)]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Arial'; st.font.size = Pt(sz); st.font.bold = True; st.font.color.rgb = col

# US Letter, 1" margins
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
CW = Inches(6.5)

# ── helpers ──────────────────────────────────────────────────────────────────
def para(text='', size=11, bold=False, italic=False, color=None, align=None, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align: p.alignment = align
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def step(n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f'{n}. ').bold = True
    p.add_run(text)
    return p

def step_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.color.rgb = NAVY
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, widths_in, bold=False, white=False, fill=None):
    cell.width = Inches(widths_in)
    if fill: shade(cell, fill)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(10); r.bold = bold
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else DARK

def make_table(headers, rows, widths_in):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, widths_in[i], bold=True, white=True, fill=HEADFILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            first_bold = (i == 0)
            set_cell(cells[i], val, widths_in[i], bold=first_bold,
                     fill=(LIGHT if i == 0 else None))
    return t

def deep_dive(title, by, manual, auto, diff):
    doc.add_heading(title, level=2)
    para(f'Automated by: {by}', italic=True, color=GREY, after=4)
    step_label('What the manager did manually')
    for i, s in enumerate(manual, 1): step(i, s)
    step_label('What the system does now')
    for i, s in enumerate(auto, 1): step(i, s)
    step_label('Key difference')
    para(diff)

# ── Title page ──────────────────────────────────────────────────────────────
para('', after=120)
para('RTA Monitoring System', size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
para('From Manual Process to Agentic Automation', size=16, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
para('How RTA managers worked before — and how the system does it now',
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
para('17 June 2026', size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── Contents ──────────────────────────────────────────────────────────────────
para('Contents', size=14, bold=True, color=NAVY, after=10)
for t in ['1. Executive Summary', '2. The Three Manual Workflows (Before)',
          '3. How the System Automates Each', '4. Workflow Deep Dive',
          '5. Extra Capabilities', '6. Architecture Overview', '7. Benefits Summary']:
    para(t, size=11, color=DARK, after=4)
doc.add_page_break()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
doc.add_heading('1. Executive Summary', level=1)
para('Real-Time Analyst (RTA) managers monitor contact-centre Service Levels (SLA) across 10 regions and 60 skills. Until now this was a manual, screen-by-screen process: managers watched live dashboards, cross-checked the CMS agent screen, took screenshots, studied spreadsheets, and hand-wrote alerts and lever reports into Microsoft Teams and email.')
para('That work was slow, repetitive, dependent on someone being at their desk, and prone to inconsistency — and any breach occurring outside staffed hours could be missed entirely.')
para('The RTA Agentic System automates this end to end. Four cooperating agents continuously scrape the dashboards and CMS, decide what action is needed, and send the same alerts and reports the managers used to produce by hand — within seconds, 24/7, and in a consistent format. This document maps each manual workflow to the part of the system that now performs it.')

# ── 2. Manual workflows ────────────────────────────────────────────────────────
doc.add_heading('2. The Three Manual Workflows (Before)', level=1)
para('Based on the RTA managers’ own process, the day-to-day monitoring work fell into three recurring workflows.')

doc.add_heading('Workflow 1 — Real-time SLA monitoring & agent recall', level=2)
for i, s in enumerate([
    'Watch each skill’s Service Level (SL) on the RTA dashboard.',
    'When a skill’s SL drops, open the CMS agent screen for that skill.',
    'Identify agents sitting on AUX (break, lunch, case management) who could return to calls.',
    'Take a screenshot of the CMS screen.',
    'Post a message in Teams asking those agents to come back on calls.'], 1): step(i, s)

doc.add_heading('Workflow 2 — SLA breach “Lever” reporting (90 → 80 → 70)', level=2)
for i, s in enumerate([
    'Notice a skill cross a lever threshold (SL falls below 90%, 80%, or 70%).',
    'Identify the skill name and its combined queue.',
    'Open the Voice_Queue intraday Excel sheet and find that queue’s interval data.',
    'Take a screenshot of the relevant rows.',
    'Analyse the root cause — high offer, high AHT, low availability, and so on.',
    'Write and email a “lever” report to the management group.'], 1): step(i, s)

doc.add_heading('Workflow 3 — AUX / AHT monitoring', level=2)
for i, s in enumerate([
    'Periodically open CMS and review each agent’s AUX state and time.',
    'Cross-check who has exceeded their break, lunch, or AHT targets.',
    'Post a Teams alert flagging those agents.'], 1): step(i, s)

# ── 3. Mapping ──────────────────────────────────────────────────────────────────
doc.add_heading('3. How the System Automates Each', level=1)
para('Each manual workflow now maps directly to one or more agents in the system:')
make_table(
    ['Manual workflow', 'Automated by', 'What the system does'],
    [['1 — SLA & agent recall', 'Agent 1 + Agent 2', 'Agent 1 scrapes SL every 60s; Agent 2 reads CMS, ranks who to recall, and posts the Teams alert automatically.'],
     ['2 — Lever reporting', 'Agent 3', 'Reads the Excel queue data, an LLM writes root cause / callouts / mitigation, and a formatted lever email is sent.'],
     ['3 — AUX / AHT monitoring', 'Agent 4', 'Continuously checks AUX / AHT / ACW against thresholds and posts a Teams alert when any is exceeded.']],
    [1.7, 1.6, 3.2])

# ── 4. Deep dive ──────────────────────────────────────────────────────────────
doc.add_heading('4. Workflow Deep Dive', level=1)
deep_dive('Workflow 1 — SLA monitoring & agent recall', 'Agent 1 (Scraper) + Agent 2 (Analyst)',
    ['Watch each skill’s SL on the dashboard.',
     'Open the CMS screen for the affected skill.',
     'Spot AUX agents who could return to calls.',
     'Screenshot CMS and post a recall request in Teams.'],
    ['Agent 1 scrapes every region’s dashboard every 60 seconds, recording SL, queue, oldest-call-waiting and availability for each skill.',
     'When a skill is breached, Agent 2 automatically opens that skill’s CMS data in a headless browser and reads every agent’s state and AUX time.',
     'Agent 2 ranks AUX agents by configurable priority and time-on-AUX and selects who to recall.',
     'It posts a structured Teams alert naming the skill, its SL, and the specific agents to bring back — no screenshot required.'],
    'The manager pasted a CMS screenshot into Teams; the system instead sends a structured text alert that names each agent and the reason. It is searchable, consistent, and produced in seconds.')
deep_dive('Workflow 2 — Lever reporting', 'Agent 3 (Lever Generator)',
    ['Spot a skill crossing 90% / 80% / 70%.',
     'Map the skill to its combined queue.',
     'Open the Voice_Queue Excel and screenshot the rows.',
     'Analyse root cause and email a lever report.'],
    ['Agent 1 detects the threshold crossing (90 / 80 / 70 = Amber / Red / Black lever).',
     'Agent 3 maps the skill to its combined queue and reads the Voice_Queue intraday Excel automatically.',
     'An LLM analyses the interval data and writes the root cause, business callouts, and mitigation actions.',
     'Agent 3 sends a fully formatted lever email: header banner, real-time snapshot, full KPI table (SL, Offered, Handled, AHT, AUX 0–9%, and more), and the SL lever-threshold legend.'],
    'Manual Excel screenshotting and hand-written analysis are replaced by an automatically generated, consistently formatted report. The lever fires once per threshold crossing and resets when the skill recovers.')
deep_dive('Workflow 3 — AUX / AHT monitoring', 'Agent 4 (CMS Monitor)',
    ['Periodically open CMS and review AUX states/times.',
     'Cross-check who exceeded break / lunch / AHT targets.',
     'Post a Teams alert flagging those agents.'],
    ['Agent 4 polls CMS every 60 seconds for every region and skill.',
     'It checks each agent’s AUX-code time against configured limits (e.g. Break ≤ 15 min, Lunch ≤ 30 min) plus AHT and ACW targets.',
     'When an agent exceeds a threshold it posts a Teams alert; the same data is shown live on the dashboard’s AUX Thresholds page.'],
    'Continuous automatic monitoring replaces periodic manual checks — nothing is missed in the gaps between a manager’s reviews.')

# ── 5. Extra capabilities ──────────────────────────────────────────────────────
doc.add_heading('5. Extra Capabilities', level=1)
doc.add_heading('Chatbot', level=3)
para('Managers can ask questions in plain English. The chatbot automatically decides whether to answer from live data or the 7-day history — for example “Who is on AUX in India?” (live) or “What was the SLA for TS_VICHW yesterday?” (history) — and tags each answer with its data source.')
doc.add_heading('Datastore', level=3)
para('A rolling 7-day history database stores every poll of every skill, powering historical questions, trends, and comparisons that were previously impossible without manual record-keeping.')
doc.add_heading('Excel integration', level=3)
para('The Voice_Queue intraday workbook remains the source of truth for lever analysis; Agent 3 reads it automatically so no one has to open it by hand.')

# ── 6. Architecture ─────────────────────────────────────────────────────────────
doc.add_heading('6. Architecture Overview', level=1)
para('The system follows a simple Scrape → Analyse → Act pipeline:')
bullet('Agent 1 (Scraper) collects live metrics from each region’s dashboard every 60 seconds.')
bullet('A router dispatches each skill to Agent 2 (real-time recall alerts) and Agent 3 (lever reports) as needed.')
bullet('Agent 4 (CMS Monitor) runs on its own 60-second loop for AUX / AHT / ACW monitoring.')
bullet('Shared state files coordinate everything: live_state.json (current snapshot for the dashboard and chatbot), cms_agents.json (current CMS agents for the AUX view), and history.db (the 7-day history store).')
bullet('Alerts are delivered to Microsoft Teams (webhooks) and managers’ inboxes (email).')
para('Coverage: 10 regions, 60 skills, refreshed every 60 seconds.', bold=True)

# ── 7. Benefits ─────────────────────────────────────────────────────────────────
doc.add_heading('7. Benefits Summary', level=1)
make_table(
    ['Benefit', 'What it means'],
    [['Speed', 'Alerts and reports are produced in seconds, versus minutes of manual screen-hopping per event.'],
     ['24/7 coverage', 'Monitoring runs every 60 seconds around the clock — no breach is missed outside staffed hours.'],
     ['Consistency', 'Every alert and lever report uses the same logic and format, regardless of who is on shift.'],
     ['No screenshots', 'Structured, searchable text alerts replace pasted CMS / Excel screenshots.'],
     ['Audit trail', 'A 7-day history of every metric supports trends, comparisons, and after-the-fact review.'],
     ['Focus shift', 'Managers move from repetitive monitoring to handling the exceptions that genuinely need judgement.']],
    [1.8, 4.7])

out = sys.argv[1] if len(sys.argv) > 1 else os.path.join(os.path.dirname(__file__), '..', 'documents', 'RTA_Manual_vs_Automation.docx')
doc.save(out)
print('Wrote', out)
